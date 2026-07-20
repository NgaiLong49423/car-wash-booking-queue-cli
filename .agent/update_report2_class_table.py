from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SOURCE = Path("docs/reports/Report_2_AutoWash_Priority_Booking_Engine.docx")
OUTPUT = Path("docs/reports/Report_2_AutoWash_Priority_Booking_Engine_completed.docx")

ROWS = [
    ("Customer", "id, name, phone, membershipLevel, points, totalSpent, visitCount",
     "getters/setters", "Managed by CustomerService; referenced by Vehicle and Booking through customerId"),
    ("Vehicle", "id, licensePlate, customerId", "getters/setters",
     "Belongs to a Customer; managed by VehicleService; referenced by Booking"),
    ("WashPackage", "id, name, price, duration, status", "getters/setters",
     "Managed by WashServiceManager; selected in Booking through serviceId"),
    ("Booking", "bookingId, customerId, vehicleId, serviceId, date, period, bookingStatus, paymentStatus, paymentMethod, createdTime",
     "getters/setters", "Managed by BookingService; enters the main queue or priority waitlist"),
    ("WaitlistEntry", "booking, tierPriority", "getBooking(), compareTo()",
     "Wraps a Booking; stored in MyPriorityQueue<WaitlistEntry>"),
    ("CompletionRecord", "completedBooking, promotedBooking", "getters",
     "Stored in MyStack<CompletionRecord> to support undo"),
    ("Period", "date, periodName, status", "getters/setters, isActive()",
     "Managed by SimulationService; defines the active booking period"),
    ("History", "bookingId, customerId, customerName, plateNumber, serviceName, completedTime, amountPaid, loyaltyPointsEarned",
     "getters/setters", "Completion history stored in MyLinkedList<History>"),
    ("Result classes\n(BookingActionResult, CancellationResult, CompletionResult, UndoResult, PeriodActivationResult)",
     "successful, message, and operation-specific result data", "failure() where applicable; getters",
     "Returned by booking, cancellation, completion, undo, and period-activation operations"),
    ("CustomerService", "customerList", "addCustomer(), findCustomerById(), searchCustomers(), updateCustomer(), deleteCustomer()",
     "Uses MyLinkedList<Customer>"),
    ("VehicleService", "vehicleList", "addVehicle(), findVehicleByLicense(), findVehicleById(), findVehicleByIdOrLicense()",
     "Uses MyLinkedList<Vehicle>; works with CustomerService"),
    ("WashServiceManager", "serviceList", "findServiceById(), isServiceActive(), sortServicesByPrice(), sortServicesByDuration()",
     "Uses MyLinkedList<WashPackage>"),
    ("BookingService", "bookingQueue, bookingList, completionStack, waitlist, bookingWindows",
     "createBooking(), activatePeriod(), processNextBooking(), confirmPayment(), addToWaitlist(), pollHighestPriorityWaitlist()",
     "Uses MyQueue, MyLinkedList, MyPriorityQueue, MyStack, and MyMap"),
    ("CompletionService", "bookingService, customerService, washServiceManager, vehicleService, historyList",
     "completeBooking(), recalculateLoyalty()", "Completes bookings, records history, and promotes waitlist bookings"),
    ("CancellationService", "bookingService, washServiceManager", "cancelAsCustomer(), cancelAsAdmin()",
     "Cancels bookings and may promote the highest-priority waitlist booking"),
    ("UndoService", "bookingService, completionService, customerService, historyList", "undoLastCompletion()",
     "Restores the most recent completion using CompletionRecord"),
    ("SimulationService", "periods", "displayCurrentTime(), setCurrentDate(), setCurrentPeriod(), activateCurrentPeriod()",
     "Uses MyLinkedList<Period> and BookingService"),
    ("HistoryService", "none", "displayGlobalHistory(), displayCustomerHistory()",
     "Reads MyLinkedList<History> for reports"),
    ("Node<T>", "data, next", "constructor", "Building block of MyLinkedList, MyQueue, and MyStack"),
    ("MyLinkedList<T>", "head, tail, size", "addLast(), addFirst(), get(), set(), remove(), clear()",
     "Stores customers, vehicles, packages, bookings, periods, and history"),
    ("MyQueue<T>", "front, rear, size", "enqueue(), dequeue(), peek(), clear()",
     "Main first-come-first-served booking queue; uses Node<T>"),
    ("MyPriorityQueue<T>", "heap, size, capacity", "insert(), poll(), peek(), snapshotInPriorityOrder()",
     "Heap-based priority waitlist; stores WaitlistEntry"),
    ("MyStack<T>", "top, size", "push(), pop(), peek(), clear()",
     "Stores CompletionRecord for undo; uses Node<T>"),
    ("MyMap<K, V>", "entries, size, capacity", "put(), get(), remove(), containsKey(), clear()",
     "Stores booking-window configuration by membership level"),
]


def set_cell_margins(cell, top=60, start=80, bottom=60, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_cell_text(cell, text, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 0.95
    for index, line in enumerate(text.split("\n")):
        if index:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = "Arial"
        run.font.size = Pt(7.5)
        run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def set_table_geometry(table, widths):
    table.autofit = False
    table.allow_autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(int(w * 1440) for w in widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(int(width * 1440)))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            tc_w = cell._tc.tcPr.tcW
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def main():
    document = Document(SOURCE)
    table = document.tables[1]

    # Retain only the header row and replace all prior descriptions.
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    headers = ("Class Name", "Attributes", "Key Methods", "Relationship")
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])

    for row_values in ROWS:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[index], value, alignment=alignment)

    set_table_geometry(table, (1.05, 1.75, 2.00, 2.10))
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
