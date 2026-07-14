from pathlib import Path
from shutil import copy2
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(r"D:\Semester 4\CSD201\autowash-priority-booking-engine")
TEMPLATE = Path(r"C:\Users\Lenovo\Downloads\Report_1_Template.docx")
OUT = ROOT / "docs" / "reports" / "Report_1_AutoWash_Priority_Booking_Engine.docx"
ASSETS = ROOT / ".report1-work" / "assets"
ASSETS.mkdir(parents=True, exist_ok=True)


def font(size, bold=False):
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
    ]
    for name in candidates:
        if Path(name).exists():
            return ImageFont.truetype(name, size)
    return ImageFont.load_default()


def box(draw, xy, text, fill, outline="#1F4E79", size=24, bold=False):
    x1, y1, x2, y2 = xy
    draw.rectangle(xy, fill=fill, outline=outline, width=3)
    f = font(size, bold)
    words = text.split()
    lines, current = [], ""
    for word in words:
        candidate = (current + " " + word).strip()
        if draw.textbbox((0, 0), candidate, font=f)[2] <= (x2 - x1 - 28):
            current = candidate
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    heights = [draw.textbbox((0, 0), line, font=f)[3] for line in lines]
    y = (y1 + y2 - sum(heights) - 8 * (len(lines) - 1)) / 2
    for line, h in zip(lines, heights):
        w = draw.textbbox((0, 0), line, font=f)[2]
        draw.text(((x1 + x2 - w) / 2, y), line, fill="#17365D", font=f)
        y += h + 8


def arrow(draw, start, end, color="#365F91", width=4):
    draw.line([start, end], fill=color, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        pts = [(ex, ey), (ex - 16 if ex > sx else ex + 16, ey - 9), (ex - 16 if ex > sx else ex + 16, ey + 9)]
    else:
        pts = [(ex, ey), (ex - 9, ey - 16 if ey > sy else ey + 16), (ex + 9, ey - 16 if ey > sy else ey + 16)]
    draw.polygon(pts, fill=color)


def curve(draw, p0, p1, p2, color, width=6):
    points = []
    for i in range(41):
        t = i / 40
        x = (1 - t) ** 2 * p0[0] + 2 * (1 - t) * t * p1[0] + t ** 2 * p2[0]
        y = (1 - t) ** 2 * p0[1] + 2 * (1 - t) * t * p1[1] + t ** 2 * p2[1]
        points.append((x, y))
    draw.line(points, fill=color, width=width)


def mind_map(path):
    im = Image.new("RGB", (1600, 1000), "white")
    d = ImageDraw.Draw(im)
    center = (625, 390, 975, 610)
    d.ellipse(center, fill="#F3F7FB", outline="#17365D", width=4)
    central = "AutoWash\nPriority Booking\nEngine"
    f = font(34, True)
    y = 432
    for line in central.split("\n"):
        w = d.textbbox((0, 0), line, font=f)[2]
        d.text(((1600 - w) / 2, y), line, fill="#17365D", font=f)
        y += 45
    branches = [
        # side, end y, colour, branch label, two detail labels
        ("left", 185, "#1F77B4", "Customer Management", "Customers", "Vehicles"),
        ("left", 490, "#70AD47", "Booking Management", "Create Booking", "Validate Tier / Window"),
        ("left", 790, "#ED7D31", "Service Management", "Wash Packages", "Search / Sort"),
        ("right", 185, "#A5007D", "Queue Management", "Main Queue (FIFO)", "Waitlist (Priority)"),
        ("right", 490, "#2F2FAD", "Service Operations", "Payment", "Completion / Cancellation"),
        ("right", 790, "#C00000", "Data & Recovery", "History / Loyalty", "Undo (Stack)"),
    ]
    for side, ey, color, label, detail1, detail2 in branches:
        if side == "left":
            start, bend, end = (630, 500), (485, ey), (295, ey)
            curve(d, start, bend, end, color)
            d.text((315, ey - 42), label, fill="#202020", font=font(29, False))
            curve(d, end, (205, ey - 42), (75, ey - 42), color, 4)
            curve(d, end, (205, ey + 42), (75, ey + 42), color, 4)
            d.text((80, ey - 82), detail1, fill="#202020", font=font(25, False))
            d.text((80, ey + 48), detail2, fill="#202020", font=font(25, False))
        else:
            start, bend, end = (970, 500), (1115, ey), (1305, ey)
            curve(d, start, bend, end, color)
            d.text((1015, ey - 64), label, fill="#202020", font=font(29, False))
            curve(d, end, (1395, ey - 42), (1525, ey - 42), color, 4)
            curve(d, end, (1395, ey + 42), (1525, ey + 42), color, 4)
            d.text((1260, ey - 82), detail1, fill="#202020", font=font(25, False))
            d.text((1260, ey + 48), detail2, fill="#202020", font=font(25, False))
    im.save(path)


def flowchart(path):
    im = Image.new("RGB", (1800, 1300), "white")
    d = ImageDraw.Draw(im)
    title = "Booking Creation and Allocation Flow"
    f = font(34, True)
    tw = d.textbbox((0, 0), title, font=f)[2]
    d.text(((1800 - tw) / 2, 28), title, fill="#17365D", font=f)
    # Top row: left to right. Only start/end are ellipses; all other steps are rectangles.
    d.ellipse((45, 260, 205, 350), fill="#EAF2F8", outline="#1F4E79", width=3)
    d.text((94, 289), "Start", fill="#202020", font=font(23, False))
    top = [
        ((275, 225, 535, 385), "Enter booking\ndetails"),
        ((605, 225, 865, 385), "Validate\ninput data"),
        ((935, 225, 1195, 385), "Check booking\nwindow"),
        ((1265, 225, 1525, 385), "Check period\ncapacity"),
    ]
    for coords, text in top:
        box(d, coords, text, "#F7F9FB", size=32, bold=False)
    arrow(d, (205, 305), (275, 305))
    for (_, _, x2, _), (x1, _, _, _) in zip([x[0] for x in top], [x[0] for x in top][1:]):
        arrow(d, (x2, 305), (x1, 305))
    # Failure paths are labelled split paths, not decision diamonds.
    failures = [
        ((735, 385), (735, 480), "No", "Validation error"),
        ((1065, 385), (1065, 480), "No", "Booking-window error"),
        ((1395, 385), (1395, 480), "No", "Period is full"),
    ]
    for start, end, label, text in failures:
        arrow(d, start, end, "#C00000", 3)
        d.text((start[0] + 10, 455), label, fill="#C00000", font=font(20, False))
        box(d, (end[0] - 125, 480, end[0] + 125, 565), text, "#FDE9E7", outline="#C00000", size=27, bold=False)
        arrow(d, (end[0], 565), (end[0], 605), "#C00000", 3)
        d.ellipse((end[0] - 65, 605, end[0] + 65, 680), fill="#FDE9E7", outline="#C00000", width=3)
        d.text((end[0] - 25, 631), "End", fill="#202020", font=font(18, False))
    d.text((875, 280), "Yes", fill="#2E7D32", font=font(18, False))
    d.text((1205, 280), "Yes", fill="#2E7D32", font=font(18, False))
    d.text((1535, 280), "Yes", fill="#2E7D32", font=font(18, False))
    # Continue on a second left-to-right row. The continuation line runs below
    # the main flow so it does not cross any process boxes.
    d.line([(1525, 305), (1650, 305), (1650, 775), (200, 775), (200, 900)], fill="#365F91", width=4)
    bottom = [
        ((250, 825, 510, 975), "Create\nWAITING booking"),
        ((590, 825, 850, 975), "Save booking\nto file"),
        ((930, 825, 1190, 975), "Current period\nactive?"),
        ((1270, 825, 1530, 975), "Main Queue /\nPriority Waitlist"),
    ]
    for coords, text in bottom:
        box(d, coords, text, "#F7F9FB", size=31, bold=False)
    for (_, _, x2, _), (x1, _, _, _) in zip([x[0] for x in bottom], [x[0] for x in bottom][1:]):
        arrow(d, (x2, 900), (x1, 900))
    arrow(d, (200, 900), (250, 900))
    d.text((1200, 860), "Yes", fill="#2E7D32", font=font(18, False))
    # The 'No' branch is a labelled lower branch; it does not overlap an error path.
    arrow(d, (1060, 975), (1060, 1065), "#365F91", 3)
    d.text((1075, 1015), "No", fill="#365F91", font=font(18, False))
    box(d, (930, 1065, 1190, 1145), "Future booking", "#F7F9FB", size=25, bold=False)
    arrow(d, (1060, 1145), (1060, 1185), "#365F91", 3)
    d.ellipse((995, 1185, 1125, 1260), fill="#EAF2F8", outline="#1F4E79", width=3)
    d.text((1030, 1211), "End", fill="#202020", font=font(20, False))
    arrow(d, (1530, 900), (1660, 900))
    d.ellipse((1660, 855, 1785, 945), fill="#EAF2F8", outline="#1F4E79", width=3)
    d.text((1695, 887), "End", fill="#202020", font=font(20, False))
    im.save(path)


def replace_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        paragraph.add_run(text)


def set_label_body(paragraph, label, body):
    replace_text(paragraph, "")
    label_run = paragraph.add_run(label)
    label_run.bold = True
    label_run.underline = False
    body_run = paragraph.add_run(body)
    body_run.bold = False
    body_run.underline = False


def find_para(doc, starts):
    for p in doc.paragraphs:
        if p.text.strip().startswith(starts):
            return p
    raise ValueError(starts)


def set_cell(cell, text, bold=False):
    p = cell.paragraphs[0]
    replace_text(p, text)
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.name = "Helvetica Neue"
        r.font.size = Pt(9)
        r.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def replace_placeholder_paragraph(doc, marker, image_path, width):
    p = find_para(doc, marker)
    replace_text(p, "")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width))


def main():
    mind = ASSETS / "mind_map.png"
    flow = ASSETS / "booking_flowchart.png"
    mind_map(mind)
    flowchart(flow)
    copy2(TEMPLATE, OUT)
    doc = Document(OUT)

    # Cover and student information
    replace_text(find_para(doc, "<PROJECT-NAME>"), "AUTOWASH PRIORITY BOOKING ENGINE")
    table = doc.tables[0]
    cover = [
        ("Full Name", "Ngô Gia Long - SE190732"),
        ("Full Name", "Nguyễn Anh Kiệt - SE190095"),
        ("Full Name", "Ngô Hoàng Thái Dương - SE190177"),
        ("Full Name", "Nguyễn Tuấn Minh - SE204693"),
        ("Instructor", "Huỳnh Công Việt Ngữ"),
        ("Report Number", "Report 1"),
        ("Class", "SE2019"),
        ("Project Topic", "AutoWash Priority Booking Engine"),
    ]
    while len(table.rows) < len(cover) + 1:
        table.add_row()
    for i, (field, value) in enumerate(cover, start=1):
        set_cell(table.cell(i, 0), field, bold=True)
        set_cell(table.cell(i, 1), value)

    # Case study overview
    set_label_body(find_para(doc, "Project Title:"), "Project Title: ", "AutoWash Priority Booking Engine")
    set_label_body(find_para(doc, "Problem Statement:"), "Problem Statement: ", "A car wash shop has limited service capacity in each morning, afternoon, and evening period. Customers may book in advance or request a slot during the active period, while higher membership tiers deserve earlier booking access and priority when capacity is limited. The system must manage customers, vehicles, wash packages, bookings, waiting queues, and completed-service history without using a database. The core data-structure challenge is to balance FIFO service order with tier-based priority and reliable rollback of the latest completion.")
    replace_text(find_para(doc, "Scope of the System:"), "Scope of the System:")
    scope = [
        "Customer, vehicle, and wash-package management with validation and file persistence.",
        "Booking creation with membership booking-window and capacity validation.",
        "Main Queue and priority Waitlist management for each service period.",
        "Payment confirmation, completion, loyalty recalculation, history, cancellation, and undo.",
    ]
    placeholders = [p for p in doc.paragraphs if p.text.strip().startswith("(Feature / functionality") or p.text.strip().startswith("(Add more as needed)")]
    for p, text in zip(placeholders, scope):
        replace_text(p, text)

    # Data structures table
    ds = doc.tables[1]
    rows = [
        ("Customer, Vehicle and Service Storage", "MyLinkedList<T>", "Stores dynamic entities and supports sequential CRUD operations without built-in collections."),
        ("Main Service Queue", "MyQueue<Booking>", "Preserves FIFO order for bookings accepted into the guaranteed service slots."),
        ("Priority Waitlist", "MyPriorityQueue<Booking>", "Selects the highest tier first; earlier-created bookings break ties."),
        ("Undo Completion", "MyStack<Booking>", "Supports LIFO rollback of the most recently completed booking."),
        ("Booking-Window Configuration", "MyMap<String, Integer>", "Maps a membership tier to its allowed number of booking days."),
    ]
    while len(ds.rows) < len(rows) + 1:
        ds.add_row()
    for i, row in enumerate(rows, start=1):
        for j, value in enumerate(row):
            set_cell(ds.cell(i, j), value)

    # Functional modules
    mods = [
        "Module 1 - Customer and Vehicle Management: maintains customer profiles, membership data, and owned vehicles.",
        "Module 2 - Service Management: manages wash packages, prices, durations, search, and sorting.",
        "Module 3 - Booking and Validation: creates bookings and checks ownership, dates, booking windows, and capacity.",
        "Module 4 - Period and Queue Management: activates a period and allocates bookings to Main Queue or Waitlist.",
        "Module 5 - Service Operations: processes, records payment, completes or cancels bookings, and promotes waitlisted customers.",
        "Module 6 - History, Loyalty and Undo: stores completed services, recalculates loyalty, and rolls back the latest completion.",
    ]
    mod_ps = [p for p in doc.paragraphs if p.text.strip().startswith("Module ") or p.text.strip().startswith("(Add more modules")]
    for p, text in zip(mod_ps, mods[:len(mod_ps)]):
        replace_text(p, text)
    anchor = mod_ps[-1]
    for text in reversed(mods[len(mod_ps):]):
        new_p = anchor.insert_paragraph_before(text)
        new_p.style = anchor.style
    # Keep the visual order stable even though Word inserts new paragraphs before the anchor.
    all_mod_ps = [p for p in doc.paragraphs if p.text.strip().startswith("Module ")]
    for p, text in zip(all_mod_ps, mods):
        replace_text(p, text)

    # Patterns table and justification
    pat = doc.tables[2]
    patterns = [
        ("FIFO Queue", "First accepted, first served processing.", "Main Queue for guaranteed service slots."),
        ("Priority Queue", "Tier-based ordering with creation time as a tie-breaker.", "Waitlist and allocation during period activation."),
        ("LIFO Stack", "The most recent operation is reversed first.", "Undo for the last completed booking."),
        ("Linear Search and Selection Sort", "Sequential lookup and comparison-based sorting.", "Entity search and wash-package sorting by price or duration."),
    ]
    for i, row in enumerate(patterns, start=1):
        for j, value in enumerate(row):
            set_cell(pat.cell(i, j), value)
    set_label_body(find_para(doc, "Justification:"), "Justification: ", "The system needs two different ordering policies. A FIFO queue keeps the guaranteed-service process predictable, while a priority queue ensures that waitlisted Platinum, Gold, Silver, and Member customers are treated according to the stated business rule. A stack is appropriate because an undo request always targets only the latest completed booking. Linear search and selection sort are suitable for this academic CLI scope and demonstrate fundamental data-structure operations.")

    # Diagrams and captions
    replace_text(find_para(doc, "(Insert your mind map"), "")
    replace_text(find_para(doc, "(Insert flowchart"), "")
    replace_placeholder_paragraph(doc, "[INSERT MIND MAP IMAGE]", mind, 6.4)
    replace_text(find_para(doc, "Caption: Mind map"), "Caption: Mind map of the AutoWash Priority Booking Engine modules and their relationships.")
    replace_placeholder_paragraph(doc, "[INSERT FLOWCHART IMAGE]", flow, 6.3)
    replace_text(find_para(doc, "Caption: Flowchart"), "Caption: Initial logic flow for creating and allocating a car-wash booking.")

    # Research question
    set_label_body(find_para(doc, "Formulated Research Question:"), "Formulated Research Question: ", "How does using a priority queue instead of a standard FIFO queue affect booking fairness and processing efficiency when a car-wash waitlist contains customers with different membership tiers?")
    set_label_body(find_para(doc, "Why this RQ matters:"), "Why this RQ matters: ", "The question directly examines the central trade-off in this project: arrival order is simple and predictable, but membership priority is necessary when capacity is scarce. Investigating the two approaches helps justify the use of a max-heap priority queue and explains how the chosen structure supports faster retrieval of the highest-priority waitlisted booking.")

    # Use regular body text. Bold is retained only for headings, table headers,
    # cover labels, and the short labels deliberately created above.
    headings = ("AUTOWASH", "REPORT 1", "STUDENT INFORMATION", "1.1", "1.2", "1.3", "1.4", "1.5", "Patterns Identified:")
    for p in doc.paragraphs:
        if p.text.strip():
            p.paragraph_format.space_after = Pt(5)
            for r in p.runs:
                if not r.font.name:
                    r.font.name = "Helvetica Neue"
                r.underline = False
                if not p.text.strip().startswith(headings):
                    r.bold = False
    # Restore intentional bolding on the labels within the explanatory paragraphs.
    for prefix in ("Project Title:", "Problem Statement:", "Justification:", "Formulated Research Question:", "Why this RQ matters:"):
        p = find_para(doc, prefix)
        if p.runs:
            p.runs[0].bold = True
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
